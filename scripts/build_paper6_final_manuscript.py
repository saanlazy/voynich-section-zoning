#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path
from typing import List, Tuple, Dict

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak, KeepTogether

ROOT = Path('/Users/admin/.gemini/antigravity/codex/paper6_section_zoning_work')
DRAFT = Path('/Users/admin/Downloads/Chang 2026g Voynich Section Specific Structural Zoning Draft.docx')
ASSETS = ROOT / 'outputs/final_manuscript_assets'
OUTDIR = ROOT / 'paper'
OUTDIR.mkdir(parents=True, exist_ok=True)
MD_OUT = OUTDIR / 'Chang_2026g_Voynich_Section_Specific_Structural_Zoning_Final.md'
DOCX_OUT = OUTDIR / 'Chang_2026g_Voynich_Section_Specific_Structural_Zoning_Final.docx'
PDF_OUT = OUTDIR / 'Chang_2026g_Voynich_Section_Specific_Structural_Zoning_Final.pdf'
CHECK_OUT = OUTDIR / 'paper6_final_render_check.md'

TABLES = {
    'Table 1': ASSETS/'tables/table1_section_corpus_summary_final.csv',
    'Table 2': ASSETS/'tables/table2_pcs_metrics_by_section_final.csv',
    'Table 3': ASSETS/'tables/table3_global_vs_section_pcs_model_final.csv',
    'Table 4': ASSETS/'tables/table4_inventory_vs_ordering_decomposition_final.csv',
    'Table 5': ASSETS/'tables/table5_section_classification_results_final.csv',
    'Table 6': ASSETS/'tables/table6_robustness_checks_final.csv',
    'Table 7': ASSETS/'tables/table7_final_key_findings_final.csv',
}
TABLE_MDS = {k: v.with_suffix('.md') for k, v in TABLES.items()}
FIGURES = {
    'Figure 1': ASSETS/'figures/figure1_pcs_conditional_entropy_by_section_final.png',
    'Figure 2': ASSETS/'figures/figure2_global_vs_section_pcs_model_final.png',
    'Figure 3': ASSETS/'figures/figure3_cross_section_transfer_heatmap_final.png',
    'Figure 4': ASSETS/'figures/figure4_inventory_vs_ordering_by_section_final.png',
    'Figure 5': ASSETS/'figures/figure5_section_classification_confusion_matrix_final.png',
    'Figure 6': ASSETS/'figures/figure6_advanced_structural_zoning_model_final.png',
}

TABLE_CAPTIONS = {
    'Table 1': 'Table 1. Section Corpus Summary. Corpus-level descriptive statistics by conventional visual/codicological section label, including token counts, type counts, type-token ratio, entropy, token length, Zipf slope, and top tokens. Section labels are treated as manuscript-description labels, not semantic categories.',
    'Table 2': 'Table 2. PCS Metrics by Section. Prefix-core-suffix metrics by section, including H(suffix|core), H(core|prefix), mutual information, and suffix-from-core prediction accuracy. These values summarize token-internal structural variation across conventional manuscript regions.',
    'Table 3': 'Table 3. Global versus Section-Specific PCS Model Comparison. Held-out comparison of a single global PCS model and section-specific PCS models. Positive delta cross-entropy values indicate cases where the section-specific model fits held-out tokens better than the global model.',
    'Table 4': 'Table 4. Inventory versus Ordering Decomposition. Decomposition of section-level entropy differences into inventory contribution and ordering contribution. The table distinguishes restricted token inventories from additional local ordering constraints.',
    'Table 5': 'Table 5. Section Classification Results. Classification of conventional visual/codicological section labels from structural features and baseline models. Predictive performance is interpreted as structural correlation with labels, not semantic classification.',
    'Table 6': 'Table 6. Robustness Checks. Robustness checks for key section-specific findings under unknown-section exclusion, low-count exclusion, token-count matched downsampling, label-certainty restriction, and label permutation baselines.',
    'Table 7': 'Table 7. Final Key Findings. Summary of the principal findings, supporting experiments, interpretation, limitations, and recommended manuscript claim strength. Claim strength is limited to structural zoning and section-conditioned regimes.',
}
FIG_CAPTIONS = {
    'Figure 1': 'Figure 1. PCS Conditional Entropy by Section. Section-level comparison of H(suffix|core) and H(core|prefix). Lower H(suffix|core) indicates stronger suffix constraint given the core within a conventional visual/codicological section.',
    'Figure 2': 'Figure 2. Global versus Section-Specific PCS Model Performance. Held-out cross-entropy for global and section-specific PCS models. Lower section-specific cross-entropy supports section-conditioned token-formation regimes where improvements are stable.',
    'Figure 3': 'Figure 3. Cross-Section Transfer Heatmap. PCS transfer performance when models trained on one section are tested on another. Transfer is interpreted as structural similarity or divergence, not semantic similarity.',
    'Figure 4': 'Figure 4. Inventory versus Ordering Contribution by Section. Entropy decomposition separating token-inventory effects from local ordering effects. This distinction prevents low section entropy from being overread as sequence organization alone.',
    'Figure 5': 'Figure 5. Section Classification Confusion Matrix. Prediction of conventional visual/codicological labels from structural features. The figure evaluates structural separability, not semantic category identification.',
    'Figure 6': 'Figure 6. Advanced Structural Zoning Model Summary. Compact summary of advanced evidence for section-conditioned structural regimes, including model comparison, transfer, inventory/order decomposition, classification, and robustness.',
}

INSERT_AFTER_SUGGESTION = {
    'Suggested insertion: Table 1. Section Corpus Summary.': [('table', 'Table 1')],
    'Suggested insertion: Table 2. PCS Metrics by Section.\nSuggested insertion: Figure 1. PCS Conditional Entropy by Section.': [('table', 'Table 2'), ('figure', 'Figure 1')],
    'Suggested insertion: Table 3. Global versus Section-Specific PCS Model Comparison.\nSuggested insertion: Figure 2. Global versus Section-Specific PCS Model Performance.': [('table', 'Table 3'), ('figure', 'Figure 2')],
    'Suggested insertion: Figure 3. Cross-Section Transfer Heatmap.\nSupplementary insertion: Supplementary Table S5. Cross-Section Transfer Matrix.': [('figure', 'Figure 3')],
    'Suggested insertion: Table 4. Inventory versus Ordering Decomposition.\nSuggested insertion: Figure 4. Inventory versus Ordering Contribution by Section.': [('table', 'Table 4'), ('figure', 'Figure 4')],
    'Suggested insertion: Table 5. Section Classification Results.\nSuggested insertion: Figure 5. Section Classification Confusion Matrix.': [('table', 'Table 5'), ('figure', 'Figure 5')],
    'Suggested insertion: Table 6. Robustness Checks.\nSupplementary insertion: Supplementary Figure S6. Robustness Summary.': [('table', 'Table 6')],
}


def sanitize_text(t: str) -> str:
    return t.replace('\u00a0', ' ').strip()


def read_draft_blocks() -> List[Tuple[str, str]]:
    doc = Document(DRAFT)
    blocks = []
    skip = False
    for p in doc.paragraphs:
        text = sanitize_text(p.text)
        if not text:
            continue
        if text == 'Tables and Figures to Insert':
            skip = True
            continue
        if skip:
            continue
        style = p.style.name
        # Drop suggested/supplementary insertion paragraphs; actual assets are inserted instead.
        if text.startswith('Suggested insertion:') or text.startswith('Supplementary insertion:'):
            blocks.append(('insertions', text))
            continue
        blocks.append((style, text))
        # Insert final summary figure and table after final conclusion paragraph.
        if text.startswith('Together, the findings support a structural zoning interpretation.'):
            blocks.append(('insertions', 'FINAL_SUMMARY_ASSETS'))
    return blocks


def format_num(v):
    if pd.isna(v):
        return 'not available'
    if isinstance(v, float):
        if abs(v) >= 1000:
            return f'{v:.2f}'
        return f'{v:.4f}'.rstrip('0').rstrip('.')
    return str(v)


def compact_col_name(c: str) -> str:
    repl = {
        'section_label': 'section', 'total_lines': 'lines', 'total_tokens': 'tokens', 'alphabetic_tokens': 'alpha tok.',
        'token_types': 'types', 'type_token_ratio': 'TTR', 'hapax_count': 'hapax', 'mean_token_length': 'mean len.',
        'median_token_length': 'median len.', 'token_entropy': 'token H', 'normalized_token_entropy': 'norm H',
        'zipf_slope': 'Zipf slope', 'top_10_tokens': 'top tokens', 'section_share_of_total_corpus': 'share',
        'pcs_valid_token_count': 'PCS n', 'prefix_type_count': 'prefix types', 'core_type_count': 'core types',
        'suffix_type_count': 'suffix types', 'prefix_entropy': 'prefix H', 'core_entropy': 'core H', 'suffix_entropy': 'suffix H',
        'H_suffix_given_core': 'H(suffix|core)', 'H_core_given_prefix': 'H(core|prefix)', 'MI_core_suffix': 'MI(core;suffix)',
        'MI_prefix_core': 'MI(prefix;core)', 'suffix_from_core_top1_accuracy': 'suffix top-1',
        'suffix_from_core_top3_accuracy': 'suffix top-3', 'core_from_prefix_top1_accuracy': 'core top-1',
        'core_from_prefix_top3_accuracy': 'core top-3', 'delta_global_minus_section_cross_entropy': 'delta CE',
        'global_cross_entropy': 'global CE', 'section_cross_entropy': 'section CE', 'train_token_count': 'train n',
        'test_token_count': 'test n', 'alpha_smoothing': 'alpha', 'valid_PCS_test_token_count': 'valid PCS n',
        'real_bigram_entropy': 'real bigram H', 'section_internal_shuffled_entropy': 'section shuffle H',
        'global_shuffled_entropy': 'global shuffle H', 'inventory_matched_shuffled_entropy': 'inventory matched H',
        'line_internal_shuffled_entropy': 'line shuffle H', 'inventory_contribution': 'inventory contrib.',
        'ordering_contribution': 'ordering contrib.', 'line_ordering_contribution': 'line-order contrib.',
        'ordering_to_inventory_ratio': 'ordering/inventory', 'macro_F1': 'macro-F1', 'weighted_F1': 'weighted-F1',
        'balanced_accuracy': 'balanced acc.', 'target_result': 'target', 'key_finding': 'finding',
        'strongest_section_or_contrast': 'section/contrast', 'manuscript_claim_strength': 'claim strength',
    }
    return repl.get(c, c.replace('_', ' '))


def table_df(label: str) -> pd.DataFrame:
    df = pd.read_csv(TABLES[label])
    # Keep all data, but shorten headers and values for document display.
    df = df.copy()
    df.columns = [compact_col_name(c) for c in df.columns]
    return df.fillna('not available')


def add_docx_paragraph(doc, text, style='Body Text', bold=False, italic=False, alignment=None, size=None):
    p = doc.add_paragraph(style=style if style in [s.name for s in doc.styles] else None)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if size:
        run.font.size = Pt(size)
    return p


def set_cell_text(cell, text, bold=False, font_size=6.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill='EDEDED'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_docx_table(doc, label: str):
    df = table_df(label)
    add_docx_paragraph(doc, TABLE_CAPTIONS[label], style='Body Text', bold=True, size=9)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    fs = 5.5 if len(df.columns) > 12 else 6.5
    for j, col in enumerate(df.columns):
        set_cell_text(table.rows[0].cells[j], col, bold=True, font_size=fs)
        set_cell_shading(table.rows[0].cells[j])
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            set_cell_text(cells[j], format_num(row[col]) if not isinstance(row[col], str) else row[col], font_size=fs)
    add_docx_paragraph(doc, '', style='Body Text')


def add_docx_figure(doc, label: str):
    path = FIGURES[label]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.3))
    cap = add_docx_paragraph(doc, FIG_CAPTIONS[label], style='Body Text', italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    cap.paragraph_format.space_after = Pt(8)


def build_docx(blocks):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    for style in doc.styles:
        if style.type == 1:
            try:
                style.font.name = 'Times New Roman'
                style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            except Exception:
                pass
    doc.styles['Normal'].font.size = Pt(11)
    for style, text in blocks:
        if style == 'insertions':
            for kind, label in ( [('figure', 'Figure 6'), ('table', 'Table 7')] if text == 'FINAL_SUMMARY_ASSETS' else INSERT_AFTER_SUGGESTION.get(text, []) ):
                if kind == 'table':
                    add_docx_table(doc, label)
                else:
                    add_docx_figure(doc, label)
            continue
        if style == 'Heading 1':
            p = add_docx_paragraph(doc, text, style='Title', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=16)
            p.paragraph_format.space_after = Pt(12)
        elif style == 'Heading 2':
            add_docx_paragraph(doc, text, style='Heading 1', bold=True, size=13)
        elif style == 'Heading 3':
            add_docx_paragraph(doc, text, style='Heading 2', bold=True, size=11.5)
        elif style == 'Compact':
            p = add_docx_paragraph(doc, '• ' + text, style='Body Text', size=10.5)
            p.paragraph_format.left_indent = Inches(0.25)
        else:
            add_docx_paragraph(doc, text, style='Body Text', size=11)
    doc.save(DOCX_OUT)


def build_markdown(blocks):
    lines = []
    for style, text in blocks:
        if style == 'insertions':
            inserts = [('figure', 'Figure 6'), ('table', 'Table 7')] if text == 'FINAL_SUMMARY_ASSETS' else INSERT_AFTER_SUGGESTION.get(text, [])
            for kind, label in inserts:
                if kind == 'table':
                    lines.append('\n' + TABLE_CAPTIONS[label] + '\n')
                    lines.append(TABLE_MDS[label].read_text(encoding='utf-8'))
                else:
                    rel = '../outputs/final_manuscript_assets/figures/' + FIGURES[label].name
                    lines.append(f'\n![{label}]({rel})\n')
                    lines.append(FIG_CAPTIONS[label] + '\n')
            continue
        if style == 'Heading 1':
            lines.append('# ' + text + '\n')
        elif style == 'Heading 2':
            lines.append('## ' + text + '\n')
        elif style == 'Heading 3':
            lines.append('### ' + text + '\n')
        elif style == 'Compact':
            lines.append('- ' + text)
        else:
            lines.append(text + '\n')
    MD_OUT.write_text('\n'.join(lines), encoding='utf-8')


def para_style_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CenterTitle', parent=styles['Title'], fontName='Times-Bold', fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name='H2x', parent=styles['Heading1'], fontName='Times-Bold', fontSize=13, leading=16, spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name='H3x', parent=styles['Heading2'], fontName='Times-Bold', fontSize=11, leading=14, spaceBefore=8, spaceAfter=4))
    styles.add(ParagraphStyle(name='Bodyx', parent=styles['BodyText'], fontName='Times-Roman', fontSize=9.5, leading=12, spaceAfter=5))
    styles.add(ParagraphStyle(name='Captionx', parent=styles['BodyText'], fontName='Times-Italic', fontSize=8.5, leading=10, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name='TableCapx', parent=styles['BodyText'], fontName='Times-Bold', fontSize=8.5, leading=10, spaceBefore=6, spaceAfter=4))
    return styles


def rl_para(text, style):
    esc = (text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
    return Paragraph(esc, style)


def rl_table(label, styles):
    df = table_df(label)
    # For PDF readability, use all rows but cap very verbose cell text.
    data = [[rl_para(str(c), styles['Bodyx']) for c in df.columns]]
    for _, row in df.iterrows():
        rr = []
        for c in df.columns:
            val = row[c]
            s = format_num(val) if not isinstance(val, str) else val
            if len(s) > 80:
                s = s[:77] + '...'
            rr.append(rl_para(s, styles['Bodyx']))
        data.append(rr)
    page_width = landscape(letter)[0] - 0.7*inch
    col_width = page_width / len(df.columns)
    t = Table(data, colWidths=[col_width] * len(df.columns), repeatRows=1)
    fs = 5 if len(df.columns) > 12 else 6
    t.setStyle(TableStyle([
        ('FONT', (0,0), (-1,-1), 'Times-Roman', fs),
        ('FONT', (0,0), (-1,0), 'Times-Bold', fs),
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('GRID', (0,0), (-1,-1), 0.25, colors.grey),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 2), ('RIGHTPADDING', (0,0), (-1,-1), 2),
    ]))
    return [PageBreak(), rl_para(TABLE_CAPTIONS[label], styles['TableCapx']), t, Spacer(1, 8)]


def rl_figure(label, styles):
    path = FIGURES[label]
    with Image.open(path) as im:
        w, h = im.size
    maxw = 6.5 * inch
    maxh = 4.8 * inch
    scale = min(maxw / w, maxh / h)
    return [Spacer(1, 6), RLImage(str(path), width=w*scale, height=h*scale), rl_para(FIG_CAPTIONS[label], styles['Captionx'])]


def build_pdf(blocks):
    styles = para_style_pdf()
    story = []
    for style, text in blocks:
        if style == 'insertions':
            inserts = [('figure', 'Figure 6'), ('table', 'Table 7')] if text == 'FINAL_SUMMARY_ASSETS' else INSERT_AFTER_SUGGESTION.get(text, [])
            for kind, label in inserts:
                story.extend(rl_table(label, styles) if kind == 'table' else rl_figure(label, styles))
            continue
        if style == 'Heading 1':
            story.append(rl_para(text, styles['CenterTitle']))
        elif style == 'Heading 2':
            story.append(rl_para(text, styles['H2x']))
        elif style == 'Heading 3':
            story.append(rl_para(text, styles['H3x']))
        elif style == 'Compact':
            story.append(rl_para('• ' + text, styles['Bodyx']))
        else:
            story.append(rl_para(text, styles['Bodyx']))
    doc = SimpleDocTemplate(str(PDF_OUT), pagesize=letter, rightMargin=0.45*inch, leftMargin=0.45*inch, topMargin=0.55*inch, bottomMargin=0.55*inch)
    doc.build(story)


def validate_outputs():
    doc_text = '\n'.join(p.text for p in Document(DOCX_OUT).paragraphs)
    md_text = MD_OUT.read_text(encoding='utf-8')
    lower = doc_text.lower()
    # Guardrail check targets affirmative overclaims, while allowing conservative
    # negated statements such as "does not propose translation" or
    # "not a deciphered stem".
    overclaim_patterns = [
        'has been deciphered',
        'is deciphered',
        'has been translated',
        'is translated',
        'source language identified',
        'confirmed syntax',
    ]
    negated_contexts = [
        'does not propose decipherment, translation, source-language identification, or confirmed syntax',
        'not a deciphered stem',
        'does not imply that the core has been translated',
        'independent of decipherment, translation, source-language identification, or confirmed syntax',
    ]
    guarded_lower = lower
    for phrase in negated_contexts:
        guarded_lower = guarded_lower.replace(phrase, '')
    checks = {
        'docx_exists': DOCX_OUT.exists(), 'markdown_exists': MD_OUT.exists(), 'pdf_exists': PDF_OUT.exists(),
        'all_tables_in_docx': all(f'Table {i}.' in doc_text for i in range(1,8)),
        'all_figures_in_docx': all(f'Figure {i}.' in doc_text for i in range(1,7)),
        'no_suggested_insertions_docx': 'Suggested insertion:' not in doc_text,
        'no_management_section_docx': 'Tables and Figures to Insert' not in doc_text,
        'no_forbidden_claims_docx': not any(x in guarded_lower for x in overclaim_patterns),
        'all_tables_in_markdown': all(f'Table {i}.' in md_text for i in range(1,8)),
        'all_figures_in_markdown': all(f'Figure {i}.' in md_text for i in range(1,7)),
    }
    lines = ['# Paper 6 Final Render Check', '', '## Output Files', f'- Markdown: `{MD_OUT}`', f'- DOCX: `{DOCX_OUT}`', f'- PDF: `{PDF_OUT}`', '', '## Structural Checks']
    lines += [f'- {k}: {v}' for k, v in checks.items()]
    lines.append('\n## Visual Render QA')
    lines.append('LibreOffice/soffice is not available in this environment, so DOCX page-image rendering could not be performed. The DOCX was structurally validated by text extraction, and a separate PDF was generated directly with ReportLab.')
    CHECK_OUT.write_text('\n'.join(lines) + '\n', encoding='utf-8')
    if not all(checks.values()):
        raise SystemExit('Validation failed: ' + str(checks))


def main():
    blocks = read_draft_blocks()
    build_markdown(blocks)
    build_docx(blocks)
    build_pdf(blocks)
    validate_outputs()
    print({'markdown': str(MD_OUT), 'docx': str(DOCX_OUT), 'pdf': str(PDF_OUT), 'check': str(CHECK_OUT)})

if __name__ == '__main__':
    main()
